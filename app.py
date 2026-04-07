from flask import Flask, render_template, request, redirect, url_for
import random
import pytesseract
from PIL import Image
import os
import re
import pdfplumber
from docx import Document
from docx.shared import Pt

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = Flask(__name__)

# ---------- CREATE FOLDERS ----------
os.makedirs("uploads", exist_ok=True)
os.makedirs("static", exist_ok=True)

# ---------- HOME ----------
@app.route("/")
def home():
    return render_template("index.html")


# ---------- SELECT INPUT ----------
@app.route("/input", methods=["GET", "POST"])
def input_page():
    if request.method == "POST":
        exam_type = request.form["exam"]
        return render_template("input_method.html", exam_type=exam_type)

    return render_template("input_method.html")


@app.route("/input_method", methods=["POST"])
def input_method():
    method = request.form["method"]

    if method == "manual":
        return render_template("input.html")

    elif method == "upload":
        return render_template("upload_questions.html")


# ---------- PROCESS UPLOAD ----------
@app.route("/process_upload", methods=["POST"])
def process_upload():

    file = request.files["question_bank"]
    filename = file.filename

    file_path = os.path.join("uploads", filename)
    file.save(file_path)

    # ---------- EXTRACT TEXT ----------
    if filename.endswith(".pdf"):
        text = extract_pdf_text(file_path)

    elif filename.endswith(".docx"):
        text = extract_docx_text(file_path)

    else:
        text = extract_image_text(file_path)

    # ---------- EXTRACT QUESTIONS ----------
    units = extract_questions_by_unit(text)

    unit1 = units.get("unit1", [])
    unit2 = units.get("unit2", [])

    if len(unit1) < 4 or len(unit2) < 4:
        return "Not enough questions"

    # ---------- SELECT QUESTIONS ----------
    selected_u1 = random.sample(unit1, 4)
    selected_u2 = random.sample(unit2, 4)

    # ---------- GENERATE DOCX ----------
    output_path = generate_docx_paper(selected_u1, selected_u2)

    return render_template("output.html", url=output_path)


# ---------- MANUAL GENERATION ----------
@app.route("/generate_manual", methods=["POST"])
def generate_manual():

    unit1 = request.form.getlist("unit1")
    unit2 = request.form.getlist("unit2")

    if len(unit1) < 4 or len(unit2) < 4:
        return "Enter at least 4 questions for each unit"

    selected_u1 = random.sample(unit1, 4)
    selected_u2 = random.sample(unit2, 4)

    template_file = request.files["template"]

    template_path = os.path.join("static", template_file.filename)
    template_file.save(template_path)

    output_path = generate_docx_manual(
        selected_u1,
        selected_u2,
        template_path
    )

    # ✅ FIXED (use UI instead of plain HTML)
    return render_template("output.html", url=output_path)


# ---------- EXTRACTION FUNCTIONS ----------
def extract_pdf_text(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() + "\n"
    return text


def extract_docx_text(file_path):
    doc = Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text


def extract_image_text(file_path):
    image = Image.open(file_path)
    return pytesseract.image_to_string(image)


# ---------- QUESTION EXTRACTION ----------
def extract_questions_by_unit(text):

    units = {}
    text = text.replace("\r", "\n")

    unit_pattern = r'(UNIT\s*\d+)'
    matches = list(re.finditer(unit_pattern, text, re.IGNORECASE))

    for i, match in enumerate(matches):

        unit_number = re.search(r'\d+', match.group()).group()

        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)

        unit_text = text[start:end]

        questions = re.findall(r'\d+[\.\)]\s*(.*)', unit_text)

        units[f"unit{unit_number}"] = [
            q.strip() for q in questions if len(q.strip()) > 5
        ]

    return units


# ---------- DOCX GENERATION ----------
def generate_docx_paper(u1, u2):

    doc = Document("static/template.docx")

    questions = u1 + u2
    q_index = 0

    for i, para in enumerate(doc.paragraphs):

        text = para.text.strip()

        if text.startswith("Q.1 A"):
            para.clear()
            run = para.add_run(f"Q.1 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.1"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.2 A"):
            para.clear()
            run = para.add_run(f"Q.2 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.2"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.3 A"):
            para.clear()
            run = para.add_run(f"Q.3 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.3"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.4 A"):
            para.clear()
            run = para.add_run(f"Q.4 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.4"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        para.paragraph_format.space_after = Pt(10)

    output_path = "static/generated_paper.docx"
    doc.save(output_path)

    return output_path


# ---------- MANUAL DOCX ----------
def generate_docx_manual(u1, u2, template_path):

    doc = Document(template_path)

    questions = u1 + u2
    q_index = 0

    for i, para in enumerate(doc.paragraphs):

        text = para.text.strip()

        if text.startswith("Q.1 A"):
            para.clear()
            run = para.add_run(f"Q.1 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.1"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.2 A"):
            para.clear()
            run = para.add_run(f"Q.2 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.2"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.3 A"):
            para.clear()
            run = para.add_run(f"Q.3 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.3"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text.startswith("Q.4 A"):
            para.clear()
            run = para.add_run(f"Q.4 A  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        elif text == "B" and doc.paragraphs[i-1].text.startswith("Q.4"):
            para.clear()
            run = para.add_run(f"B  {questions[q_index]}")
            run.bold = True
            run.font.size = Pt(12)
            q_index += 1

        para.paragraph_format.space_after = Pt(10)

    output_path = "static/manual_generated_paper.docx"
    doc.save(output_path)

    return output_path

# ---------- TEMPLATES PAGE ----------
@app.route("/templates")
def templates():
    return render_template("templates.html")


# ---------- HISTORY PAGE ----------
@app.route("/history")
def history():
    files = os.listdir("static")
    return render_template("history.html", files=files)

# ---------- RUN ----------
if __name__ == "__main__":
   import os
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 10000))
    app.run(host="0.0.0.0", port=port)